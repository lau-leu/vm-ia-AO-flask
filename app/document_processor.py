"""Utilitaires pour le traitement des documents Word et PDF."""
import hashlib
from pathlib import Path
from typing import Optional
import re

from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
import pdfplumber


def calculate_file_hash(file_path: str) -> str:
    """Calculer le hash SHA256 d'un fichier."""
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()


def extract_text_from_pdf(file_path: str) -> str:
    """Extraire le texte d'un fichier PDF."""
    text_parts = []

    try:
        # Essayer avec pdfplumber pour une meilleure extraction
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception:
        # Fallback vers PyPDF2
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            return f"Erreur lors de l'extraction du PDF: {str(e)}"

    return "\n\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    """Extraire le texte d'un fichier Word (.docx)."""
    try:
        doc = DocxDocument(file_path)
        text_parts = []

        # Extraire le texte des paragraphes
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extraire le texte des tableaux
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)
    except Exception as e:
        return f"Erreur lors de l'extraction du DOCX: {str(e)}"


def extract_text(file_path: str) -> str:
    """Extraire le texte d'un document (PDF ou Word)."""
    path = Path(file_path)
    extension = path.suffix.lower()

    if extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif extension in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        return f"Format de fichier non supporté: {extension}"


def create_word_document(
    content: str,
    title: str,
    reference: str,
    output_path: str,
    template_path: Optional[str] = None
) -> str:
    """
    Créer un document Word à partir du contenu généré.

    Args:
        content: Contenu textuel de l'offre
        title: Titre du document
        reference: Référence de l'offre
        output_path: Chemin de sortie
        template_path: Chemin vers un modèle Word (optionnel)

    Returns:
        Chemin du fichier créé
    """
    # Charger le template ou créer un nouveau document
    if template_path and Path(template_path).exists():
        doc = DocxDocument(template_path)
    else:
        doc = DocxDocument()

    # Ajouter l'en-tête
    header = doc.add_heading(title, level=0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajouter la référence
    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_run = ref_para.add_run(f"Référence: {reference}")
    ref_run.bold = True
    ref_run.font.size = Pt(12)

    doc.add_paragraph()  # Ligne vide

    # Parser et ajouter le contenu
    sections = parse_generated_content(content)

    for section in sections:
        if section['type'] == 'heading':
            doc.add_heading(section['text'], level=section.get('level', 1))
        elif section['type'] == 'paragraph':
            doc.add_paragraph(section['text'])
        elif section['type'] == 'list':
            for item in section['items']:
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(item)
        elif section['type'] == 'table':
            add_table_to_doc(doc, section['data'])

    # Sauvegarder le document
    doc.save(output_path)
    return output_path


def parse_generated_content(content: str) -> list:
    """
    Parser le contenu généré par l'IA pour structurer le document.

    Returns:
        Liste de sections avec leur type et contenu
    """
    sections = []
    lines = content.split('\n')
    current_list = []

    for line in lines:
        line = line.strip()
        if not line:
            if current_list:
                sections.append({'type': 'list', 'items': current_list})
                current_list = []
            continue

        # Détecter les titres (markdown style)
        if line.startswith('# '):
            if current_list:
                sections.append({'type': 'list', 'items': current_list})
                current_list = []
            sections.append({'type': 'heading', 'text': line[2:], 'level': 1})
        elif line.startswith('## '):
            if current_list:
                sections.append({'type': 'list', 'items': current_list})
                current_list = []
            sections.append({'type': 'heading', 'text': line[3:], 'level': 2})
        elif line.startswith('### '):
            if current_list:
                sections.append({'type': 'list', 'items': current_list})
                current_list = []
            sections.append({'type': 'heading', 'text': line[4:], 'level': 3})
        # Détecter les listes
        elif line.startswith('- ') or line.startswith('* '):
            current_list.append(line[2:])
        elif re.match(r'^\d+\.\s', line):
            current_list.append(re.sub(r'^\d+\.\s', '', line))
        else:
            if current_list:
                sections.append({'type': 'list', 'items': current_list})
                current_list = []
            sections.append({'type': 'paragraph', 'text': line})

    if current_list:
        sections.append({'type': 'list', 'items': current_list})

    return sections


def add_table_to_doc(doc: DocxDocument, data: list):
    """Ajouter un tableau au document Word."""
    if not data:
        return

    rows = len(data)
    cols = len(data[0]) if data else 0

    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = str(cell_text)


def extract_key_information(text: str) -> dict:
    """
    Extraire les informations clés d'un appel d'offre.

    Returns:
        Dictionnaire avec les informations extraites
    """
    info = {
        'reference': None,
        'title': None,
        'deadline': None,
        'budget': None,
        'requirements': [],
        'criteria': []
    }

    # Patterns pour extraire les informations
    patterns = {
        'reference': r'(?:référence|ref|n°)\s*[:\-]?\s*([A-Z0-9\-\/]+)',
        'deadline': r'(?:date limite|échéance|deadline)\s*[:\-]?\s*(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',
        'budget': r'(?:budget|montant)\s*[:\-]?\s*([\d\s]+(?:€|EUR|euros?))',
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            info[key] = match.group(1).strip()

    # Extraire le titre (première ligne significative)
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if lines:
        info['title'] = lines[0][:200]  # Limiter à 200 caractères

    return info
